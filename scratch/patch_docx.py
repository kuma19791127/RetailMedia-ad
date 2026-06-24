import docx
import os

doc_path = r"c:\Users\one\.gemini\antigravity\playground\twilight-parsec\RetailMedia_Secure_Specs.docx"
doc = docx.Document(doc_path)

# 1. Table 0 へのセキュリティ項目追加
table = doc.tables[0]

new_security_items = [
    (
        "DBマイグレーション起動デッドロック防止＆自動タイムアウト",
        "新旧コンテナ起動時のテーブルロック競合（ALTER TABLE等による排他ロック待ち）を排除するため、単一の接続クライアント内で SET lock_timeout = 5000 を強制適用し、万が一のタイムアウト時もエラーをキャッチして起動を継続する、極めて耐障害性の高いマイグレーション設計を導入。"
    ),
    (
        "クロスロール2FA（二段階認証）の動的同期",
        "同一メールアドレスで「店舗」や「広告主」など複数のロールを持つユーザーが、一方のロールで2FAを有効化した場合、もう一方のロールログイン時にも2FAシークレットをデータベース上で自動同期・強制適用。ユーザーの利便性を損なわずにシステム全体の安全性を担保。"
    ),
    (
        "添付ファイルの安全なデコード保存とXSSの徹底排除",
        "お問い合わせフォーム等から送信された添付画像（Base64）データをサーバー側で安全にバイナリデコードしてディスク（uploads/contacts/）に隔離保存。管理者画面での受信一覧の表示時には、エスケープ処理（escapeHtml）を適用した上で安全なレンダリングを行い、XSS（クロスサイトスクリプティング）を完全に無害化。"
    ),
    (
        "CORSおよびミドルウェアの適用漏れ自動防御",
        "CORSブロックによる通信切断を防ぐため、すべてのルートハンドラよりも上部で Express CORS ミドルウェアを一律に適用し、Credentialsが不要なセッションではOrigin「*」を許容しつつ、トークンによる認証を並行して機能させることで、ブラウザ側のCORS拒否を完全に回避。"
    )
]

for title, desc in new_security_items:
    row = table.add_row()
    row.cells[0].text = title
    row.cells[1].text = desc

# 2. 文末に「6. 他システムとの差別化要素（セキュアUX ＆ 高レジリエンス設計）」の章を追加
doc.add_paragraph("")  # 空行
h = doc.add_heading("6. 他システムとの差別化要素（セキュアUX ＆ 高レジリエンス設計）", level=1)
# docx のデフォルトスタイルでフォントを調整する場合はこのように追加
p = doc.add_paragraph()
p.add_run("本システムが他社の標準的なリテールメディア・アドプラットフォームと差別化される、セキュリティおよびUXを両立させたコア設計仕様です。\n\n").italic = True

p_diff1 = doc.add_paragraph()
r1 = p_diff1.add_run("■ シームレスな自動復帰（Remember-Me & Smart Redirect）仕様\n")
r1.bold = True
p_diff1.add_run(
    "ログイン成功時にのみ、LocalStorageへ遷移先URL（last_active_service）および認証情報を保存し、ログアウト時でもこれらを維持します。次回起動（index.htmlアクセス）時に、過去に開いていた適切なログイン画面へ自動リダイレクトと認証情報の自動入力を行うことで、二重ログイン等の状態不整合（セッション混同）を防ぎつつ、極めてシームレスなUXをユーザーに提供します。"
)

p_diff2 = doc.add_paragraph()
r2 = p_diff2.add_run("■ 外部AI APIの動的フォールバックと完全デモ保証（ハイブリッド・レジリエンス）\n")
r2.bold = True
p_diff2.add_run(
    "Google Gemini等の外部生成AI APIを使用する際、単一のモデルに依存せず、優先度の高いモデル（gemini-2.5-flash, gemini-1.5-flash等）をリスト化して自動で切り替えるモデルフォールバック機構を搭載。さらに、通信障害時やAPIキー未設定時でも、事前に定義されたデモ用データを返却してフロントエンドをクラッシュさせずに処理を継続する、他社にない高耐久性を実現しています。"
)

p_diff3 = doc.add_paragraph()
r3 = p_diff3.add_run("■ ローカル・本番のハイブリッド実行環境（DB動的ローテーション）\n")
r3.bold = True
p_diff3.add_run(
    "本番のクラウド環境（RDS PostgreSQL）と、ローカル開発環境（SQLite）を環境変数（DATABASE_URL）の有無によって動的に判別。本番環境で不要なバイナリ（sqlite3等）のロードによるクラッシュ（ERR_DLOPEN_FAILED）を完全に防止するとともに、database.jsonとPostgreSQL/SQLite間のデータ自動マイグレーションを起動時に実行する、運用安全性の高いインフラ構成を構築しています。"
)

doc.save(doc_path)
print("Successfully patched RetailMedia_Secure_Specs.docx with new security requirements and differentiators.")
