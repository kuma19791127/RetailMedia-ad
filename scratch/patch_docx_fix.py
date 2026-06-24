import docx

doc_path = r"c:\Users\one\.gemini\antigravity\playground\twilight-parsec\RetailMedia_Secure_Specs.docx"
doc = docx.Document(doc_path)

modified = False
target_text = "database.jsonとPostgreSQL/SQLite間のデータ自動マイグレーションを起動時に実行する"
replacement_text = "インメモリJSONへの依存を完全に排除し、堅牢なリレーショナルDBへの直接SQL参照設計で動作する"

for p in doc.paragraphs:
    if target_text in p.text:
        # docx の段落のテキストを直接置換すると、装飾（runs）が失われる可能性がありますが、
        # 今回の段落は add_paragraph 時に単純なテキストとして追加されたものなので、直接 text を上書きして問題ありません。
        p.text = p.text.replace(target_text, replacement_text)
        modified = True
        print("Successfully replaced text in paragraph.")

if modified:
    doc.save(doc_path)
    print("Document saved successfully.")
else:
    print("Warning: target_text not found in any paragraph.")
