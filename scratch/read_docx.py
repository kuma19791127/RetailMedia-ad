import docx
import sys
import io

# Set stdout to UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc_path = r"c:\Users\one\.gemini\antigravity\playground\twilight-parsec\RetailMedia_Secure_Specs.docx"
doc = docx.Document(doc_path)

print(f"Number of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")

print("\n--- Paragraphs ---")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

print("\n--- Tables ---")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        row_txt = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        # Remove duplicates from merged cells
        unique_row_txt = []
        for val in row_txt:
            if not unique_row_txt or unique_row_txt[-1] != val:
                unique_row_txt.append(val)
        print(f"Row {r_idx}: {unique_row_txt}")
